# ================= 勇者使用指南 =================
# 在cmd或者powershell中运行以下指令，安装勇者必需的装备（依赖项）：
#   pip install biopython openai tqdm python-docx
# 注意：爱丽丝还没有测试过这些装备是否齐全，如果运行时报错，很可能是装备没装完，可以根据报错信息自行补充哦！
# 启动勇者程序：
#   python gene_review_bot_qwen_v2.py --file genes.txt （使用脚本里默认的物种和性状，也可以自己调整～）
# 或者指定物种和性状：
#   python gene_review_bot_qwen_v2.py --file genes.txt --species pig chicken --traits "meat quality" "growth rate" --cleanup （这样会删除临时文件，像勇者打扫战场一样！）
import os
import time
import argparse
from typing import List, Dict
import re
from Bio import Entrez, Medline
from openai import OpenAI
from tqdm import tqdm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn  # 用于设置中文字体，让勇者的文字更清晰！
import json  # 用于结构化保存文献数据，像勇者的冒险日志
import shutil  # 用于清理临时文件，打完怪要整理背包嘛～

# ================= 勇者装备配置 =================
# 【重要】请在这里填入您的API密钥，就像勇者的魔法钥匙！
DEEPSEEK_API_KEY = "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxx" 
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
# 这里记得把DEEPSEEK_API_KEY换成自己申请到的deepseek密钥哦～如果能开个公共账号大家一起用就更好啦，不然爱丽丝的零花钱会不够用的(>_<)
NCBI_API_KEY = "xxxxxxxxxxxxxxxxxxxxxxxxxxx"  # 这里要换成自己申请到的NCBI API密钥，百度一下就知道怎么拿啦！

# NCBI的勇者设置（必须配置）
Entrez.email = "your_email@example.com" 
Entrez.api_key = NCBI_API_KEY  # 启用API Key认证，就像勇者的身份徽章
Entrez.tool = "GeneReviewBot"  # 爱丽丝给工具起的名字！

# 邮箱好像可以随便填？不过最好用学校邮箱，每个人都不一样才安全呢～

# 访问速度控制（有了API Key就可以加速啦）
REQUEST_DELAY = 0.35  # 秒：NCBI建议≥0.33秒/请求（有API Key时）；其实NCBI说一秒可以处理10个请求，设置0.1秒也ok，但爱丽丝觉得0.2秒更稳妥～
BATCH_DELAY = 1.5     # 批次间的休息时间，勇者也要喘口气嘛

# 下面就是AI自动生成的备注啦

# 全局临时目录配置
TEMP_ROOT_DIR = "temp_gene_analysis"  # 临时文件都藏在这个勇者背包里！

# ================= 辅助小技能 =================
def sanitize_filename(filename: str) -> str:
    """去掉文件名里的非法字符，就像清理路上的小怪物"""
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

def set_chinese_font(paragraph):
    """给段落里的中文字体穿上‘宋体’盔甲，让它们更帅气！"""
    for run in paragraph.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ================= 勇者伙伴：类定义 =================

class PubMedFetcher:
    """这个伙伴负责去PubMed迷宫寻找文献宝藏！"""
    def __init__(self, species_list: List[str], traits_list: List[str]):
        self.species = species_list
        self.traits = traits_list
        self.current_query = None
        self.current_gene = None
        
        # 物种名称的MeSH映射表，就像物种翻译官
        self.mesh_mapping = {
            "pig": "Swine", "chicken": "Chickens", "cow": "Cattle",
            "sheep": "Sheep", "goat": "Goats", "duck": "Ducks",
            "rabbit": "Rabbits", "horse": "Horses", "turkey": "Turkeys"
        }
        
        # 确保临时根目录存在
        os.makedirs(TEMP_ROOT_DIR, exist_ok=True)

    def construct_query(self, gene: str) -> str:
        """构建精准的PubMed查询语句，像勇者施放魔法咒语"""
        species_terms = []
        for s in self.species:
            s_lower = s.lower()
            mesh_term = self.mesh_mapping.get(s_lower, s)
            species_terms.append(f'("{mesh_term}"[MeSH Terms] OR "{s}"[Title/Abstract])')
        
        species_part = " OR ".join(species_terms)
        traits_part = " OR ".join([f'"{t}"[Title/Abstract]' for t in self.traits]) if self.traits else ""
        
        query = f'("{gene}"[Gene Name] OR "{gene}"[Title/Abstract]) AND ({species_part})'
        if traits_part:
            query += f" AND ({traits_part})"
        
        self.current_query = query
        self.current_gene = gene
        return query

    def search_and_fetch(self, gene: str, max_results: int = 100) -> List[Dict[str, str]]:
        """使用NCBI API高效获取文献，就像勇者派出信鸽收集情报"""
        query = self.construct_query(gene)
        print(f"\n🔍 搜索基因: {gene}")
        print(f"   PubMed查询咒语: {query}")
        
        # 创建安全的基因文件夹名
        safe_gene_name = sanitize_filename(gene)
        # 在临时根目录下为每个基因建一个子文件夹
        temp_dir = os.path.join(TEMP_ROOT_DIR, safe_gene_name)
        os.makedirs(temp_dir, exist_ok=True)
        
        # 保存查询元数据（JSON格式，方便核对）
        meta = {
            "gene": gene,
            "query": query,
            "species": self.species,
            "traits": self.traits,
            "max_results": max_results,
            "timestamp": time.strftime('%Y-%m-%d %H:%M:%S'),
            "ncbi_api_used": bool(NCBI_API_KEY)
        }
        with open(os.path.join(temp_dir, "00_metadata.json"), "w", encoding="utf-8") as f:
            json.dump(meta, f, indent=2, ensure_ascii=False)
        
        try:
            # 第1步：搜索获取PMID列表
            search_handle = Entrez.esearch(
                db="pubmed", 
                term=query, 
                retmax=max_results, 
                sort="relevance",
                usehistory="y"  # 启用历史会话，可以一次拿很多数据！
            )
            search_results = Entrez.read(search_handle)
            search_handle.close()
            id_list = search_results["IdList"]
            webenv = search_results.get("WebEnv", "")
            query_key = search_results.get("QueryKey", "")
            
            print(f"✅ 找到 {len(id_list)} 篇文献 (PMID示例: {id_list[:3] if id_list else '无'})")
            
            if not id_list:
                with open(os.path.join(temp_dir, "01_no_results.txt"), "w", encoding="utf-8") as f:
                    f.write("这次没有找到文献，下次再挑战吧！")
                return []
            
            # 保存PMID列表（这是重要的线索哦）
            with open(os.path.join(temp_dir, "01_pmid_list.txt"), "w", encoding="utf-8") as f:
                f.write(f"总共 {len(id_list)} 个PMID:\n")
                f.write("\n".join(id_list))
            
            # 第2步：高效获取文献详情（用WebEnv防止URL过长）
            papers = []
            batch_size = 50  # 有API Key就可以一次拿更多啦！
            
            with open(os.path.join(temp_dir, "02_raw_papers.json"), "w", encoding="utf-8") as json_out, \
                 open(os.path.join(temp_dir, "02_raw_papers.txt"), "w", encoding="utf-8") as txt_out:
                
                txt_out.write(f"= 基因 {gene} 的文献原始数据（勇者搜集到的宝藏） =\n\n")
                all_records = []
                
                for start in tqdm(range(0, len(id_list), batch_size), desc="📥 获取文献详情"):
                    try:
                        fetch_handle = Entrez.efetch(
                            db="pubmed",
                            rettype="medline",
                            retmode="text",
                            retstart=start,
                            retmax=batch_size,
                            webenv=webenv,
                            query_key=query_key
                        )
                        records = list(Medline.parse(fetch_handle))
                        fetch_handle.close()
                        
                        for record in records:
                            pmid = record.get("PMID", "N/A")
                            title = record.get("TI", "").strip()
                            abstract = record.get("AB", "").strip()
                            
                            if title and abstract:  # 只有带标题和摘要的文献才算完整宝藏
                                paper = {
                                    "pmid": pmid,
                                    "title": title,
                                    "abstract": abstract,
                                    "authors": record.get("AU", []),
                                    "journal": record.get("JT", ""),
                                    "year": record.get("DP", "")[:4] if record.get("DP") else ""
                                }
                                papers.append(paper)
                                all_records.append(paper)
                                
                                # TXT格式方便勇者直接阅读
                                txt_out.write(f"PMID: {pmid}\n")
                                txt_out.write(f"标题: {title}\n")
                                txt_out.write(f"期刊: {paper['journal']} ({paper['year']})\n")
                                txt_out.write(f"摘要: {abstract}\n")
                                txt_out.write("-" * 80 + "\n\n")
                        
                        time.sleep(REQUEST_DELAY)  # 遵守NCBI的规则，慢慢来～
                    
                    except Exception as e:
                        print(f"⚠️  获取批次 {start}-{start+batch_size} 时出错: {str(e)[:100]}")
                        time.sleep(BATCH_DELAY * 2)
                
                # 保存结构化JSON，方便程序验证
                json.dump(all_records, json_out, indent=2, ensure_ascii=False)
            
            print(f"✅ 成功获取 {len(papers)} 篇有效文献（都带标题和摘要哦！）")
            return papers

        except Exception as e:
            error_msg = f"❌ PubMed搜索失败: {type(e).__name__}: {str(e)}"
            print(error_msg)
            with open(os.path.join(temp_dir, "99_error.log"), "w", encoding="utf-8") as f:
                f.write(error_msg)
            return []

class DeepSeekSummarizer:
    """这个伙伴负责用DeepSeek的智慧总结文献，就像勇者的贤者导师！"""
    def __init__(self, api_key, base_url):
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.model = "deepseek-reasoner"
        self.current_gene = None

    def summarize_batch(self, gene: str, papers: List[Dict[str, str]], batch_idx: int) -> str:
        """第一阶段：分批总结（爱丽丝会严格约束AI，不准它瞎编！）"""
        self.current_gene = gene
        safe_gene_name = sanitize_filename(gene)
        temp_dir = os.path.join(TEMP_ROOT_DIR, safe_gene_name)
        
        # 构建严格的提示词，就像给AI下达勇者任务
        paper_texts = []
        for i, p in enumerate(papers, 1):
            paper_texts.append(
                f"[文献{batch_idx*10 + i}] PMID:{p['pmid']}\n"
                f"标题: {p['title']}\n"
                f"摘要: {p['abstract']}\n"
            )
        
        prompt = (
            f"【任务】严格基于以下{len(papers)}篇文献摘要，总结基因'{gene}'在畜牧物种中的功能。\n"
            f"【铁律】\n"
            f"1. 只准用下面给的文献内容，不许自己瞎编，也不许用外面的知识！\n"
            f"2. 每个结论后面必须标注来源：(文献X) 或 (PMID:XXXXXX)\n"
            f"3. 如果文献没提到某个性状，必须写“该文献未提及[性状]”\n"
            f"4. 不准用“研究表明”“文献指出”这种模糊的说法！\n"
            f"5. 如果摘要里没有明确结论，直接说“摘要未提供具体结论”\n\n"
            f"【文献数据】\n" + "\n".join(paper_texts) + "\n"
            f"【输出要求】\n"
            f"- 用中文分点总结\n"
            f"- 每点都要带文献索引\n"
            f"- 没信息的地方要明确写“未提及”"
        )
        
        # 保存AI的输入（带PMID，方便验证）
        input_file = os.path.join(temp_dir, f"03_ai_input_batch{batch_idx+1}.txt")
        with open(input_file, "w", encoding="utf-8") as f:
            f.write(f"基因: {gene} | 批次: {batch_idx+1}\n")
            f.write(f"提交时间: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("="*80 + "\n\n")
            f.write(prompt)
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,  # 把温度调低，让AI更听话！
                max_tokens=1500
            )
            summary = response.choices[0].message.content
            
            # 保存AI的原始输出
            output_file = os.path.join(temp_dir, f"04_ai_output_batch{batch_idx+1}.txt")
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(f"基因: {gene} | 批次: {batch_idx+1}\n")
                f.write(f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"模型: {self.model}\n")
                f.write("="*80 + "\n\n")
                f.write(summary)
            
            return summary
            
        except Exception as e:
            error_msg = f"AI总结批次{batch_idx+1}失败: {str(e)}"
            print(f"⚠️  {error_msg}")
            return f"[错误] {error_msg}"

    def final_review(self, gene: str, batch_summaries: List[str]) -> str:
        """第二阶段：整合验证（爱丽丝要检查引用是否真实，防止AI做梦！）"""
        safe_gene_name = sanitize_filename(gene)
        temp_dir = os.path.join(TEMP_ROOT_DIR, safe_gene_name)
        
        # 保存所有批次总结
        with open(os.path.join(temp_dir, "05_all_batch_summaries.txt"), "w", encoding="utf-8") as f:
            for i, summary in enumerate(batch_summaries, 1):
                f.write(f"=== 批次 {i} ===\n{summary}\n\n")
        
        # 构建最终提示（强调验证）
        prompt = (
            f"【任务】整合以下关于基因'{gene}'的分批次总结，生成最终综述报告。\n"
            f"【核心要求】\n"
            f"1. 严格只用下面提供的总结内容，不准新增信息！\n"
            f"2. 每个结论必须保留原始文献索引（比如(文献3)或(PMID:123456)）\n"
            f"3. 报告结构：\n"
            f"   - 基因功能概述（基于文献）\n"
            f"   - 对生产性状的影响（肉质/生长/繁殖）\n"
            f"   - 环境适应性作用（热应激等）\n"
            f"   - 【关键】验证声明：列出所有被引用的PMID，并说明“以上结论均基于临时文件中记录的文献”\n"
            f"4. 如果某个领域没有文献支持，要明确写“当前检索文献中未发现相关证据”\n\n"
            f"【分批次总结内容】\n" + "\n---批次分隔---\n".join(batch_summaries)
        )
        
        # 保存最终整合输入
        with open(os.path.join(temp_dir, "06_final_input.txt"), "w", encoding="utf-8") as f:
            f.write(prompt)
        
        try:
            print(f"🧠 生成 {gene} 最终报告（爱丽丝正在检查引用...）")
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=2000
            )
            final_text = response.choices[0].message.content
            
            # 保存最终输出
            with open(os.path.join(temp_dir, "07_final_output.txt"), "w", encoding="utf-8") as f:
                f.write(final_text)
            
            # 生成验证报告（超重要！）
            self._generate_verification_report(gene, final_text, temp_dir)
            
            return final_text
            
        except Exception as e:
            print(f"❌ 最终综述生成失败: {str(e)}")
            return f"[错误] 最终综述生成失败: {str(e)}"

    def _generate_verification_report(self, gene: str, final_text: str, temp_dir: str):
        """生成引用验证报告（这是对抗AI幻想的大招！）"""
        import re
        cited_pmids = set(re.findall(r'PMID[:\s]*(\d+)', final_text))
        cited_indices = set(re.findall(r'文献(\d+)', final_text))
        
        # 读取实际获取的PMID
        raw_papers_path = os.path.join(temp_dir, "02_raw_papers.json")
        actual_pmids = set()
        if os.path.exists(raw_papers_path):
            with open(raw_papers_path, "r", encoding="utf-8") as f:
                papers = json.load(f)
                actual_pmids = {p["pmid"] for p in papers}
        
        # 生成验证报告
        report = []
        report.append("="*60)
        report.append(f"🔬 基因 {gene} - 引用验证报告（勇者爱丽丝的检查结果）")
        report.append(f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        report.append("="*60)
        report.append(f"\n✅ 实际获取文献PMID数量: {len(actual_pmids)}")
        report.append(f"📄 AI报告中引用的PMID数量: {len(cited_pmids)}")
        
        if cited_pmids:
            report.append("\n🔍 引用PMID验证:")
            for pmid in sorted(cited_pmids):
                status = "✓ 存在于原始文献" if pmid in actual_pmids else "✗ 未在原始文献中找到！"
                report.append(f"   PMID:{pmid} → {status}")
        
        # 检查文献索引
        if cited_indices:
            max_index = len([f for f in os.listdir(temp_dir) if f.startswith("04_ai_output_batch")])
            report.append("\n📌 文献索引验证:")
            for idx in sorted(cited_indices, key=int):
                status = "✓ 索引有效" if 1 <= int(idx) <= max_index*10 else "✗ 索引超出范围！"
                report.append(f"   文献{idx} → {status}")
        
        # 最终结论
        report.append("\n" + "="*60)
        if not cited_pmids or all(pmid in actual_pmids for pmid in cited_pmids):
            report.append("✅ 验证结论: 未发现虚构文献引用，勇者可以放心！")
        else:
            report.append("⚠️  验证结论: 发现可疑引用！请重点核查标记为'✗'的PMID")
        report.append("="*60)
        report.append("\n💡 使用建议:")
        report.append(f"1. 打开 {temp_dir}/02_raw_papers.txt 核对原始文献")
        report.append("2. 在PubMed官网搜索可疑PMID验证文献真实性")
        report.append("3. 检查AI输出中是否包含未在原始文献中出现的结论")
        
        # 保存验证报告
        report_path = os.path.join(temp_dir, "99_verification_report.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write("\n".join(report))
        
        # 控制台高亮提示
        if any("✗" in line for line in report):
            print("\n" + "\033[1;31m" + "="*70)
            print("⚠️  警告：验证报告发现可疑引用！请立即查看:")
            print(f"   {report_path}")
            print("="*70 + "\033[0m\n")
        else:
            print(f"✅ 验证通过！详细报告: {report_path}")

# ================= 辅助函数（含中文字体支持） =================
def create_master_doc(filename):
    """创建主文档并设置中文字体，让报告看起来更整洁～"""
    doc = Document()
    
    # 设置全局默认字体为宋体
    styles = doc.styles
    font_styles = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']
    
    for style_name in font_styles:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = '宋体'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    heading = doc.add_heading('基因功能综述汇总报告（勇者爱丽丝的冒险成果）', 0)
    set_chinese_font(heading)
    
    # 添加时间信息
    p = doc.add_paragraph(f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    set_chinese_font(p)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    doc.save(filename)
    return doc

def append_to_doc(filename, gene, content):
    """追加内容到Word文档，确保中文字体支持，就像勇者不断更新冒险日志"""
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = create_master_doc(filename)
    
    # 添加基因标题
    gene_heading = doc.add_heading(f"基因: {gene}", level=1)
    set_chinese_font(gene_heading)
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 处理标题
        if line.startswith('### '):
            heading = doc.add_heading(line.replace('### ', ''), level=3)
            set_chinese_font(heading)
        elif line.startswith('## '):
            heading = doc.add_heading(line.replace('## ', ''), level=2)
            set_chinese_font(heading)
        elif line.startswith('# '):
            heading = doc.add_heading(line.replace('# ', ''), level=1)
            set_chinese_font(heading)
        # 处理列表项
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            set_chinese_font(p)
        # 普通段落
        else:
            p = doc.add_paragraph(line)
            set_chinese_font(p)
    
    # 保存文档（UTF-8编码）
    doc.save(filename)

# ================= 勇者主程序 =================
def main():
    parser = argparse.ArgumentParser(description="基因文献综述勇者机器人 (NCBI API增强版)")
    parser.add_argument("--file", type=str, required=True, help="基因列表txt文件路径")
    parser.add_argument("--species", type=str, nargs='+', 
                       default=["pig", "chicken", "cow", "sheep", "goat"],
                       help="物种列表（可以选5-10个）")
    parser.add_argument("--traits", type=str, nargs='+',
                       default=["heat stress", "meat quality", "growth", "adaptation"],
                       help="性状关键词列表")
    parser.add_argument("--max-results", type=int, default=100, help="每个基因最多找多少篇文献")
    parser.add_argument("--cleanup", action="store_true", help="完成后删除临时文件，清理战场")
    
    args = parser.parse_args()
    
    # 检查API配置
    if NCBI_API_KEY == "your-ncbi-api-key-here" or not NCBI_API_KEY.strip():
        print("\033[1;31m❌ 错误：请在程序开头配置您的NCBI_API_KEY！\033[0m")
        print("获取方法：https://account.ncbi.nlm.nih.gov/settings/")
        return
    if DEEPSEEK_API_KEY == "sk-your-deepseek-api-key" or not DEEPSEEK_API_KEY.strip():
        print("\033[1;31m❌ 错误：请配置DEEPSEEK_API_KEY！\033[0m")
        return
    
    # 读取基因列表
    if not os.path.exists(args.file):
        print(f"\033[1;31m❌ 错误：文件不存在: {args.file}\033[0m")
        return
    with open(args.file, 'r', encoding='utf-8') as f:
        genes = [line.strip() for line in f if line.strip()]
    print(f"\n📋 共 {len(genes)} 个基因待处理: {', '.join(genes[:5])}{'...' if len(genes)>5 else ''}")
    
    # 确保临时根目录存在
    os.makedirs(TEMP_ROOT_DIR, exist_ok=True)
    
    # 初始化
    master_doc = "All_Genes_Summary_Report.docx"
    create_master_doc(master_doc)
    print(f"📄 主报告文件: {master_doc}")
    
    fetcher = PubMedFetcher(args.species, args.traits)
    summarizer = DeepSeekSummarizer(DEEPSEEK_API_KEY, DEEPSEEK_BASE_URL)
    
    # 开始勇者之旅，逐个基因挑战！
    for idx, gene in enumerate(genes, 1):
        print(f"\n{'='*70}")
        print(f"⚙️  当前进度: {idx}/{len(genes)} | 挑战基因: {gene}")
        print(f"{'='*70}")
        
        papers = fetcher.search_and_fetch(gene, max_results=args.max_results)
        
        if not papers:
            append_to_doc(master_doc, gene, "⚠️  没有找到带标题和摘要的有效文献，这次冒险落空了～")
            continue
        
        # 分批总结
        batch_summaries = []
        batch_size = 10
        
        print(f"\n🧠 开始AI分批总结 ({len(papers)} 篇文献 → {-(len(papers)//-batch_size)} 个批次)...")
        for i in range(0, len(papers), batch_size):
            batch = papers[i:i+batch_size]
            summary = summarizer.summarize_batch(gene, batch, i//batch_size)
            if summary and not summary.startswith("[错误]"):
                batch_summaries.append(summary)
            time.sleep(1.5)  # API调用间隔，让勇者喘口气
        
        # 生成最终报告
        if batch_summaries:
            final_report = summarizer.final_review(gene, batch_summaries)
            append_to_doc(master_doc, gene, final_report)
            
            # 保存独立报告
            safe_gene_name = sanitize_filename(gene)
            report_path = os.path.join(TEMP_ROOT_DIR, safe_gene_name, f"{safe_gene_name}_final_report.txt")
            with open(report_path, "w", encoding="utf-8") as f:
                f.write(final_report)
            
            print(f"\n✅ 基因 {gene} 处理完成！")
            print(f"   📁 临时文件目录: {os.path.join(TEMP_ROOT_DIR, safe_gene_name)}")
            print(f"   📄 独立报告: {report_path}")
        else:
            print(f"⚠️  基因 {gene} 没有生成有效总结")
            append_to_doc(master_doc, gene, "⚠️  AI总结过程出错，没能生成内容")
        
        # 基因间延迟，避免API太累
        if idx < len(genes):
            print(f"\n⏳ 等待 {BATCH_DELAY*2} 秒后挑战下一个基因...")
            time.sleep(BATCH_DELAY * 2)
    
    print("\n" + "="*70)
    print("🎉 所有基因挑战成功！勇者爱丽丝为你鼓掌！")
    print(f"📄 主汇总报告: {master_doc}")
    print(f"📂 临时文件存在: {TEMP_ROOT_DIR}/")
    print(f"🔍 验证建议: 记得检查各基因目录下的 99_verification_report.txt 哦")
    
    # 可选：清理临时文件
    if args.cleanup:
        print("\n🗑️  正在清理临时文件...")
        try:
            shutil.rmtree(TEMP_ROOT_DIR)
            print(f"✅ 临时文件目录 {TEMP_ROOT_DIR} 已删除，背包干净啦！")
        except Exception as e:
            print(f"⚠️  临时文件清理失败: {str(e)}")
    
    print("="*70)

if __name__ == "__main__":
    # 检查必要库
    try:
        from Bio import __version__ as biopython_version
        if tuple(map(int, biopython_version.split('.')[:2])) < (1, 79):
            print("\033[1;33m⚠️  警告：Biopython版本有点低（需要≥1.79），可能不支持api_key参数哦\033[0m")
    except:
        pass
    
    main()